from __future__ import annotations

import re
import time
from datetime import datetime
from typing import List, Dict, Any

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.common.exceptions import WebDriverException, NoSuchElementException, JavascriptException

from docx import Document as DocxDocument

try:
    # Optional dependency; keep alias stable with user's code
    from langchain.schema import Document as LCDocument  # type: ignore
except Exception:  # pragma: no cover - allow running without langchain installed
    class LCDocument:  # minimal shim if langchain isn't available
        def __init__(self, page_content: str, metadata: Dict[str, Any] | None = None):
            self.page_content = page_content
            self.metadata = metadata or {}


class WebCrawler:
    def __init__(self, headless: bool = True, delay: int = 2):
        """
        Args:
            headless (bool): Chạy browser ẩn
            delay (int): Thời gian nghỉ giữa các request (giây)
        """
        self.delay = delay
        self.chrome_options = Options()
        if headless:
            # Headless mới ổn định hơn
            self.chrome_options.add_argument("--headless=new")
        self.chrome_options.add_argument("--no-sandbox")
        self.chrome_options.add_argument("--disable-dev-shm-usage")
        self.chrome_options.add_argument("--disable-gpu")
        self.driver = None
        self.header = ""

    # ----------------------------- Driver ----------------------------- #
    def _setup_driver(self):
        """Khởi tạo Chrome driver"""
        try:
            self.driver = webdriver.Chrome(options=self.chrome_options)
        except WebDriverException as e:
            raise RuntimeError(f"Lỗi khởi tạo ChromeDriver: {e}")

    def _close_driver(self):
        """Đóng driver"""
        if self.driver:
            try:
                self.driver.quit()
            except Exception:
                pass
            self.driver = None

    # ----------------------------- Tables ----------------------------- #
    def _extract_table_data(self, table_element) -> List[List[str]]:
        """
        Trích xuất dữ liệu từ bảng HTML

        Args:
            table_element: WebElement của bảng

        Returns:
            list: Danh sách các hàng, mỗi hàng là list các cell
        """
        rows: List[List[str]] = []
        try:
            tr_elements = table_element.find_elements(By.TAG_NAME, "tr")
            for tr in tr_elements:
                row_data: List[str] = []
                cells = tr.find_elements(By.TAG_NAME, "th") + tr.find_elements(By.TAG_NAME, "td")
                for cell in cells:
                    # Xử lý link trong cell (hiện các link ẩn "TẠI ĐÂY" -> [Link: ...])
                    self._replace_hidden_links(cell)

                    cell_text = cell.text.strip()
                    # Thêm link [Link: href] nếu không nằm trong text
                    links = cell.find_elements(By.TAG_NAME, "a")
                    if links and not any(keyword in cell_text.upper() for keyword in ["TẠI ĐÂY", "[LINK:"]):
                        for link in links:
                            href = link.get_attribute("href")
                            if href and href not in cell_text:
                                cell_text += f" [Link: {href}]"
                    row_data.append(cell_text)
                if row_data:
                    rows.append(row_data)
        except Exception as e:
            print(f"Lỗi khi trích xuất bảng: {e}")
        return rows

    def _format_table_as_text(self, table_rows: List[List[str]]) -> str:
        """Chuyển bảng sang ASCII table (để hiển thị text song song)."""
        if not table_rows:
            return ""

        max_cols = max(len(r) for r in table_rows)
        col_widths = [0] * max_cols
        for r in table_rows:
            for i, c in enumerate(r):
                col_widths[i] = max(col_widths[i], len(str(c)))

        def sep() -> str:
            return "+" + "+".join("-" * (w + 2) for w in col_widths) + "+"

        lines = []
        for i, r in enumerate(table_rows):
            if i == 0:
                lines.append(sep())
            cells = [f" {str(c):<{col_widths[j]}} " for j, c in enumerate(r)]
            lines.append("|" + "|".join(cells) + "|")
            lines.append(sep())
        return "\n".join(lines)

    def _add_table_to_docx(self, doc: DocxDocument, table_rows: List[List[str]]) -> None:
        """Thêm bảng thực vào file DOCX."""
        if not table_rows:
            return
        try:
            max_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=max_cols)
            table.style = "Table Grid"
            for i, row_data in enumerate(table_rows):
                for j, cell_data in enumerate(row_data):
                    table.cell(i, j).text = str(cell_data)
            doc.add_paragraph("")
        except Exception as e:
            print(f"Lỗi khi thêm bảng vào docx: {e}")

    # ----------------------------- Text & Links ----------------------------- #
    def _replace_hidden_links(self, element) -> bool:
        """Hiện link thật thay cho 'TẠI ĐÂY' bên trong element."""
        try:
            links = element.find_elements(By.TAG_NAME, "a")
            for link in links:
                text_upper = (link.text or "").strip().upper()
                if any(kw in text_upper for kw in ["TẠI ĐÂY", "TẠI ĐÂY.", "TẠI ĐÂY,"]):
                    href = link.get_attribute("href")
                    if href:
                        try:
                            self.driver.execute_script(
                                "arguments[0].textContent = arguments[1];",
                                link,
                                f"[Link: {href}]",
                            )
                        except JavascriptException:
                            # Bỏ qua nếu không sửa được DOM
                            pass
            return True
        except Exception as e:
            print(f"Lỗi khi xử lý link ẩn: {e}")
            return False

    def _get_content_with_tables(self, element) -> Dict[str, Any]:
        """
        Lấy nội dung từ element, xử lý riêng text và bảng.
        Returns: {"text": str, "tables": List[List[List[str]]]}  (list bảng, mỗi bảng là list các hàng)
        """
        content: Dict[str, Any] = {"text": "", "tables": []}
        try:
            # Trước tiên, xử lý hidden links trong toàn bộ element
            self._replace_hidden_links(element)
            
            # Lấy tất cả các bảng
            tables = element.find_elements(By.TAG_NAME, "table")
            
            if tables:
                # Trích xuất dữ liệu từ các bảng
                for table in tables:
                    table_data = self._extract_table_data(table)
                    if table_data:
                        content["tables"].append(table_data)

                # Tạo bản sao của element và loại bỏ tất cả các bảng
                script = """
                    var cloned = arguments[0].cloneNode(true);
                    var tables = cloned.querySelectorAll('table');
                    for (var i = tables.length - 1; i >= 0; i--) { 
                        tables[i].parentNode.removeChild(tables[i]); 
                    }
                    return cloned.textContent || cloned.innerText || "";
                """
                try:
                    text_without_tables = self.driver.execute_script(script, element)
                    content["text"] = (text_without_tables or "").strip()
                except JavascriptException:
                    # Fallback: lấy text từ element gốc
                    content["text"] = (element.text or "").strip()
                
                # Thêm placeholder cho các bảng vào text để biết vị trí
                table_text = ""
                for i, table_data in enumerate(content["tables"]):
                    table_text += f"\n\n[BẢNG {i+1}]\n{self._format_table_as_text(table_data)}\n"
                
                content["text"] += table_text
                
            else:
                # Không có bảng, chỉ lấy text thông thường
                content["text"] = (element.text or "").strip()
                
        except Exception as e:
            print(f"Lỗi khi lấy nội dung: {e}")
            # Fallback: chỉ lấy text
            content["text"] = (element.text or "").strip()
            
        return content

    # ----------------------------- Crawl ----------------------------- #
    def _collect_urls(self, base_url: str) -> List[str]:
        """
        Thu thập tất cả URLs từ các phần tử có class 'single_content_tip'
        
        Args:
            base_url (str): URL trang chứa danh sách các single_content_tip
            
        Returns:
            List[str]: Danh sách URLs được tìm thấy
        """
        urls = []
        assert self.driver is not None, "Driver chưa được khởi tạo."
        
        try:
            print(f"Đang thu thập URLs từ: {base_url}")
            self.driver.get(base_url)
            time.sleep(3)  # Đợi trang load
            
            # Tìm tất cả elements có class single_content_tip
            tip_elements = self.driver.find_elements(By.CSS_SELECTOR, ".single_content_tip")
            print(f"Tìm thấy {len(tip_elements)} single_content_tip elements")
            
            for i, tip_element in enumerate(tip_elements):
                try:
                    # Tìm link trong mỗi tip element
                    link_selectors = ["a", "a[href]", ".tip-title a", ".content-title a"]
                    
                    for selector in link_selectors:
                        try:
                            link = tip_element.find_element(By.CSS_SELECTOR, selector)
                            href = link.get_attribute("href")
                            
                            if href and href.startswith("http"):
                                if href not in urls:  # Tránh trùng lặp
                                    urls.append(href)
                                    title = link.text.strip() or "N/A"
                                    print(f"  [{i+1}] {title}: {href}")
                                break
                        except NoSuchElementException:
                            continue
                            
                except Exception as e:
                    print(f"Lỗi khi xử lý tip element {i+1}: {e}")
                    continue
                    
        except Exception as e:
            print(f"Lỗi khi thu thập URLs: {e}")
            
        print(f"Thu thập được {len(urls)} URLs")
        return urls
    def _crawl_single_url(self, url: str) -> Dict[str, Any]:
        """Crawl nội dung từ một URL"""
        assert self.driver is not None, "Driver chưa được khởi tạo."
        try:
            self.driver.get(url)
            time.sleep(2)  # đơn giản, có thể thay bằng WebDriverWait nếu cần

            content_data: Dict[str, Any] = {"text": "", "tables": []}

            # Lấy header
            self.header = ""
            header_selectors = ['div[class="header"] h2', "h1", "h2", ".header h2"]
            for selector in header_selectors:
                try:
                    header = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if header.text.strip():
                        self.header = header.text.strip()
                        break
                except NoSuchElementException:
                    continue

            # Lấy nội dung chính
            main_selectors = [".tip-detail"] # ".main-content", ".content"
            for selector in main_selectors:
                try:
                    main_div = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if main_div.text.strip():
                        content_data = self._get_content_with_tables(main_div)
                        if content_data["text"] or content_data["tables"]:
                            break
                except NoSuchElementException:
                    continue

            if not (content_data["text"] or content_data["tables"]):
                return {"text": "Không thể lấy nội dung", "tables": []}
            return content_data
        except Exception as e:
            print(f"Lỗi khi crawl {url}: {e}")
            return {"text": f"Lỗi: {str(e)}", "tables": []}

    # ----------------------------- Output helpers ----------------------------- #
    def _clean_text_from_ascii_table(self, text: str) -> str:
        """Loại bỏ các dòng định dạng bảng ASCII (ngăn double-render trong DOCX)."""
        if not text:
            return ""
        cleaned_lines: List[str] = []
        for line in text.splitlines():
            if re.match(r"^\s*\+[-+\s]*\+\s*$", line):
                # đường viền: +-----+-----+
                continue
            if re.match(r"^\s*\|.*\|\s*$", line):
                # dòng cell: |  a  |  b  |
                continue
            cleaned_lines.append(line)
        return "\n".join(cleaned_lines).strip()

    def _create_filename(self) -> str:
        """Tạo tên file DOCX từ header nếu có."""
        clean_header = re.sub(r'[<>:"/\\\\|?*]', "", self.header or "").strip()
        return f"{clean_header}.docx" if clean_header else f"content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    def _save_to_docx(self, content_data: Dict[str, Any], filename: str) -> bool:
        """Lưu nội dung + bảng vào file DOCX"""
        try:
            doc = DocxDocument()
            if self.header:
                doc.add_heading(self.header, 0)

            text_content: str = content_data.get("text", "") or ""
            tables: List[List[List[str]]] = content_data.get("tables", []) or []

            if "[BẢNG" in text_content:
                parts = re.split(r"\[BẢNG \d+\]", text_content)
                table_index = 0
                for part in parts:
                    clean_part = self._clean_text_from_ascii_table(part)
                    if clean_part:
                        doc.add_paragraph(clean_part)
                    if table_index < len(tables):
                        self._add_table_to_docx(doc, tables[table_index])
                        table_index += 1
                while table_index < len(tables):
                    self._add_table_to_docx(doc, tables[table_index])
                    table_index += 1
            else:
                if text_content.strip():
                    doc.add_paragraph(text_content.strip())
                for table_data in tables:
                    self._add_table_to_docx(doc, table_data)

            doc.save(filename)
            return True
        except Exception as e:
            print(f"Lỗi khi lưu file {filename}: {e}")
            return False

    # ----------------------------- Public API ----------------------------- #
    def run(self, urls: List[str]) -> List[LCDocument]:
        """
        Crawl danh sách URLs và lưu thành các file DOCX

        Args:
            urls (list): Danh sách các URL cần crawl

        Returns:
            list: Danh sách LangChain Document objects
        """
        results: List[LCDocument] = []
        print(f"Bắt đầu crawl {len(urls)} trang web...")
        self._setup_driver()
        try:
            for i, url in enumerate(urls, 1):
                print(f"[{i}/{len(urls)}] Crawling: {url}")
                self.header = ""  # Reset header cho mỗi URL

                content_data = self._crawl_single_url(url)
                filename = self._create_filename()
                success = self._save_to_docx(content_data, filename)

                page_content = content_data.get("text", "")
                if content_data.get("tables"):
                    page_content += f"\n\n[Tìm thấy {len(content_data['tables'])} bảng trong nội dung]"

                result = LCDocument(
                    page_content=page_content,
                    metadata={
                        "source": url,
                        "header": self.header,
                        "tables_count": len(content_data.get("tables", [])),
                        "has_tables": bool(content_data.get("tables")),
                        "saved_as": filename,
                    },
                )
                results.append(result)

                if success:
                    print(f"✓ Đã lưu: {filename}")
                    if content_data.get("tables"):
                        print(f"  → Tìm thấy {len(content_data['tables'])} bảng")
                else:
                    print(f"✗ Lỗi khi lưu: {filename}")

                if i < len(urls):
                    time.sleep(self.delay)
        except Exception as e:
            print(f"Lỗi trong quá trình crawl: {e}")
        finally:
            self._close_driver()
            print("Đã hoàn thành crawl.")
        return results


if __name__ == "__main__":
    # Ví dụ sử dụng
    crawler = WebCrawler(headless=True, delay=3)
    crawler._setup_driver()
    urls = crawler._collect_urls("https://sv-ctt.hust.edu.vn/#/so-tay-sv") 
    print(urls)
    # docs = crawler.run(urls)

    # print("\n=== KẾT QUẢ ===")
    # for d in docs:
    #     print(f"✓ {d.metadata.get('source')}")
    #     print(f"  Header: {d.metadata.get('header', 'N/A')}")
    #     print(f"  Có bảng: {'Có' if d.metadata.get('has_tables') else 'Không'}")
    #     if d.metadata.get('tables_count', 0) > 0:
    #         print(f"  Số bảng: {d.metadata['tables_count']}")
    #     print(f"  Lưu thành: {d.metadata.get('saved_as')}")
    #     print("---")

